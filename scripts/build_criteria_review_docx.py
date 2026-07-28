# -*- coding: utf-8 -*-
"""Professional RTL Word document: evaluation criteria & final scoring review."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "مراجعة_معايير_التقييم_والنتيجة_النهائية.docx"
)
DESKTOP = Path.home() / "Desktop" / "مراجعة_معايير_التقييم_والنتيجة_النهائية.docx"

BROWN = RGBColor(92, 64, 51)
ACCENT = RGBColor(139, 69, 19)
MUTED = RGBColor(90, 80, 70)
WHITE = RGBColor(255, 255, 255)
SOFT_BG = "F7F0E4"
HEADER_BG = "8B4513"
ROW_ALT = "FBF6F0"


def _set_run_font(run, *, size: int = 12, bold: bool = False, color=BROWN) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Traditional Arabic"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Traditional Arabic")
    rfonts.set(qn("w:hAnsi"), "Traditional Arabic")
    rfonts.set(qn("w:cs"), "Traditional Arabic")
    if rpr.find(qn("w:rtl")) is None:
        rpr.append(OxmlElement("w:rtl"))


def _set_para_rtl(paragraph, *, after: int = 8, before: int = 0, align=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    paragraph.alignment = align
    ppr = paragraph._element.get_or_add_pPr()
    if ppr.find(qn("w:bidi")) is None:
        ppr.append(OxmlElement("w:bidi"))
    pf = paragraph.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_runs(paragraph, parts: list[tuple[str, dict]]) -> None:
    for text, opts in parts:
        run = paragraph.add_run(text)
        _set_run_font(
            run,
            size=int(opts.get("size", 12)),
            bold=bool(opts.get("bold", False)),
            color=opts.get("color", BROWN),
        )


def p_text(doc: Document, text: str, *, size: int = 12, bold: bool = False, after: int = 8, before: int = 0, color=BROWN) -> None:
    p = doc.add_paragraph()
    _set_para_rtl(p, after=after, before=before)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color)


def p_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_para_rtl(p, after=4, before=0)
    p.paragraph_format.right_indent = Cm(0.25)
    run = p.add_run(f"•  {text}")
    _set_run_font(run, size=12, color=BROWN)


def p_numbered(doc: Document, n: int, text: str) -> None:
    p = doc.add_paragraph()
    _set_para_rtl(p, after=4, before=0)
    p.paragraph_format.right_indent = Cm(0.25)
    run = p.add_run(f"{n}.  {text}")
    _set_run_font(run, size=12, color=BROWN)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    _set_para_rtl(p, after=10, before=16 if level == 1 else 12)
    if level == 1:
        # Accent bar via border bottom
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "18")
        bottom.set(qn("w:space"), "6")
        bottom.set(qn("w:color"), "8B4513")
        pBdr.append(bottom)
        pPr.append(pBdr)
        run = p.add_run(text)
        _set_run_font(run, size=15, bold=True, color=ACCENT)
    else:
        run = p.add_run(text)
        _set_run_font(run, size=13, bold=True, color=BROWN)


def callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _shade_cell(cell, SOFT_BG)
    _set_cell_borders(cell, "C4A990")
    p1 = cell.paragraphs[0]
    _set_para_rtl(p1, after=4, before=4)
    r1 = p1.add_run(title)
    _set_run_font(r1, size=12, bold=True, color=ACCENT)
    p2 = cell.add_paragraph()
    _set_para_rtl(p2, after=4, before=0)
    r2 = p2.add_run(body)
    _set_run_font(r2, size=11, color=MUTED)
    doc.add_paragraph()


def _shade_cell(cell, fill_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_borders(cell, color_hex: str = "C4A990") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), color_hex)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 11, color=BROWN, center: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    _set_para_rtl(
        p,
        after=2,
        before=2,
        align=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT,
    )
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _shade_cell(cell, HEADER_BG)
        _set_cell_borders(cell, HEADER_BG)
        _set_cell_text(cell, h, bold=True, size=11, color=WHITE, center=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                _shade_cell(cell, ROW_ALT)
            _set_cell_borders(cell, "D9C7B8")
            _set_cell_text(cell, val, size=11, color=BROWN)
    doc.add_paragraph()


def set_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.right_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    sect_pr = section._sectPr
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    sect_pr.append(bidi)

    # Header
    header = section.header
    hp = header.paragraphs[0]
    _set_para_rtl(hp, after=0, before=0)
    hr = hp.add_run("نظام إدارة التمارين  |  وثيقة مراجعة داخلية")
    _set_run_font(hr, size=9, color=MUTED)

    # Footer with page number field
    footer = section.footer
    fp = footer.paragraphs[0]
    _set_para_rtl(fp, after=0, before=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    fr = fp.add_run("صفحة ")
    _set_run_font(fr, size=9, color=MUTED)
    # PAGE field
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run2 = fp.add_run()
    run2._element.append(fld_begin)
    run2._element.append(instr)
    run2._element.append(fld_end)
    _set_run_font(run2, size=9, color=MUTED)


def build() -> Path:
    doc = Document()
    set_doc(doc)
    style = doc.styles["Normal"]
    style.font.name = "Traditional Arabic"
    style.font.size = Pt(12)

    # —— Cover ——
    p_text(doc, "وزارة / القوات — نظام إدارة التمارين", size=11, color=MUTED, after=4)
    p_text(doc, "مساحة المحللين", size=12, bold=True, color=ACCENT, after=18)

    title = doc.add_paragraph()
    _set_para_rtl(title, after=10, before=8)
    tr = title.add_run("آلية معايير التقييم\nواحتساب النتيجة النهائية")
    _set_run_font(tr, size=22, bold=True, color=ACCENT)

    p_text(doc, "وثيقة مراجعة مهنية للمنهجية المعتمدة في النظام", size=12, color=MUTED, after=6)
    p_text(doc, "الاتجاه: من اليمين إلى اليسار  •  للاستخدام الداخلي", size=10, color=MUTED, after=18)

    callout(
        doc,
        "الغرض من الوثيقة",
        "توثيق مسار معايير التقييم، طريقة إدخال المعايير حسب المرحلة، حساب تعدد القوائم لنفس الوحدة، ثم آلية احتساب النتيجة النهائية على مستوى المرحلة ومستوى الوحدة.",
    )

    # —— TOC-like ——
    heading(doc, "محتويات الوثيقة")
    for i, item in enumerate(
        [
            "مسار الدخول والشاشة الرئيسية",
            "أعمدة جدول توزيع النسبة المئوية",
            "خطوات إدخال المعايير حسب المرحلة",
            "اشتراطات التشغيل وتأكيد التفعيل",
            "حساب عدة قوائم لنفس الوحدة والمرحلة",
            "مثال رقمي توضيحي",
            "احتساب النتيجة النهائية",
            "جدول ملخص سريع",
            "خلاصة للمراجع",
        ],
        start=1,
    ):
        p_numbered(doc, i, item)

    # —— 1 ——
    heading(doc, "1. مسار الدخول والشاشة الرئيسية")
    p_text(doc, "يتم الوصول إلى الشاشة عبر المسار التالي:")
    p_numbered(doc, 1, "الصفحة الرئيسية ← بطاقة المحللين ← ابدأ")
    p_numbered(doc, 2, "أدوات المحللين ← معايير التقييم")
    p_text(
        doc,
        "تظهر صفحة بعنوان «توزيع النسبة المئوية الإجمالية للتقييم»، وتعرض صفوفاً لمستويات الوحدات وأعمدة للمراحل التشغيلية.",
        before=6,
    )

    # —— 2 ——
    heading(doc, "2. أعمدة جدول توزيع النسبة المئوية")
    add_table(
        doc,
        ["العمود", "المعنى التشغيلي"],
        [
            ["ت", "رقم تسلسلي للصف."],
            ["الوحدة", "مستوى الوحدة من التخطيط (مثل قيادة كتيبة، سرية إشارة، مستوى لواء)."],
            [
                "أعمدة المراحل",
                "مجموع العلامات المخصصة للوحدة في كل مرحلة (تحضير، انفتاح، معركة تعرضية، مسارات التقييم).",
            ],
            ["الإجمالي", "مجموع علامات مراحل الوحدة."],
            [
                "النسبة المئوية المخصصة",
                "(إجمالي الوحدة ÷ إجمالي كل الوحدات) × 100.",
            ],
            [
                "تفاصيل / حذف",
                "اختيار المرحلة لفتح جدول المعايير التفصيلي، أو حذف الوحدة.",
            ],
        ],
    )
    callout(
        doc,
        "دلالة الرمز «—»",
        "ظهور «—» في خلية مرحلة يعني أنه لم تُحفظ بعد علامات معايير لتلك الوحدة والمرحلة. هذا فراغ بيانات وليس تعطيلاً للآلية.",
    )

    # —— 3 ——
    heading(doc, "3. خطوات إدخال المعايير حسب المرحلة")
    p_numbered(doc, 1, "في صف الوحدة المطلوبة افتح القائمة «— اختر المرحلة —».")
    p_numbered(doc, 2, "اختر المرحلة (تحضير / انفتاح / معركة تعرضية / مسارات التقييم).")
    p_numbered(doc, 3, "ينتقل النظام إلى صفحة جدول المرحلة الخاصة بتلك الوحدة.")
    p_numbered(
        doc,
        4,
        "تظهر بنود المعايير من عناوين قوائم التقييم المنشورة في التخطيط لنفس الوحدة والمرحلة.",
    )
    p_numbered(doc, 5, "أدخل «العلامة المخصصة» لكل صف (لكل قائمة).")
    p_numbered(doc, 6, "تُحسب النسبة داخل المرحلة تلقائياً لكل صف.")
    p_numbered(doc, 7, "اضغط «حفظ وعودة» للرجوع إلى جدول التوزيع.")
    p_text(
        doc,
        "بعد الحفظ تمتلئ خلية المرحلة في الجدول العام بمجموع العلامات. يُكرَّر الإجراء لكل وحدة ولكل مرحلة حسب الحاجة.",
        before=8,
    )

    # —— 4 ——
    heading(doc, "4. اشتراطات التشغيل وتأكيد التفعيل")
    p_bullet(doc, "وجود تمرين حالي في مساحة العمل.")
    p_bullet(doc, "وجود مراحل تقييم في كتالوج التخطيط / بنك المعلومات.")
    p_bullet(
        doc,
        "وجود قوائم تقييم منشورة من التخطيط لنفس الوحدة والمرحلة؛ وإلا تظهر رسالة بعدم توفر قوائم لتلك المرحلة.",
    )
    callout(
        doc,
        "تأكيد النظام",
        "الاشتراطات أعلاه مفعّلة أصلاً في صفحة معايير التقييم. الخلايا الفارغة تعني عدم إدخال بيانات بعد، وليست غياباً للمنهجية.",
    )

    # —— 5 ——
    heading(doc, "5. حساب عدة قوائم لنفس الوحدة والمرحلة")
    p_text(
        doc,
        "إذا كان لمستوى وحدة واحد (مثل اللواء) في مرحلة واحدة عدة قوائم تقييم (مثال: 7 قوائم)، يعاملها النظام كالتالي:",
    )
    p_bullet(doc, "كل قائمة منشورة تظهر صفاً مستقلاً في جدول تفاصيل المرحلة.")
    p_bullet(doc, "تُدخل علامة مخصصة لكل قائمة على حدة.")
    p_bullet(doc, "إجمالي المرحلة لتلك الوحدة = مجموع علامات القوائم (جمع وليس متوسطاً).")
    p_bullet(doc, "نسبة كل قائمة داخل المرحلة = (علامة القائمة ÷ إجمالي المرحلة) × 100.")
    p_bullet(doc, "في جدول التوزيع العام تظهر خلية المرحلة لذلك المستوى بنفس مجموع العلامات.")
    p_text(
        doc,
        "خلاصة: القوائم المتعددة لا تُدمَج في رقم واحد قبل الإدخال؛ لكل قائمة وزنها، والحساب النهائي للمرحلة هو جمع الأوزان.",
        bold=True,
        before=8,
    )

    # —— 6 ——
    heading(doc, "6. مثال رقمي توضيحي")
    p_text(doc, "افرض علامات مخصصة لسبع قوائم في مرحلة واحدة:")
    add_table(
        doc,
        ["ت", "القائمة", "العلامة المخصصة", "النسبة داخل المرحلة"],
        [
            ["1", "قائمة 1", "10", "10%"],
            ["2", "قائمة 2", "15", "15%"],
            ["3", "قائمة 3", "10", "10%"],
            ["4", "قائمة 4", "20", "20%"],
            ["5", "قائمة 5", "15", "15%"],
            ["6", "قائمة 6", "10", "10%"],
            ["7", "قائمة 7", "20", "20%"],
            ["—", "الإجمالي", "100", "100%"],
        ],
    )
    p_text(
        doc,
        "إذا لم تُدخل علامة لقائمة معيّنة فإنها لا تدخل في المجموع حتى تُحفظ قيمة لها.",
        color=MUTED,
    )

    # —— 7 ——
    heading(doc, "7. احتساب النتيجة النهائية")
    p_text(
        doc,
        "بعد ضبط المعايير، تعتمد صفحة «التقييم نهائي» على نتائج المحكمين مع القصوى المأخوذة من معايير التقييم (ما لم تُدخل قصوى يدوية).",
    )
    p_text(doc, "الحساب يتم على مستويين معاً: المرحلة ثم مستوى الوحدة.", bold=True, before=6)

    heading(doc, "7.1 داخل المرحلة (لنفس مستوى الوحدة)", level=2)
    p_text(
        doc,
        "مثال: مستوى اللواء في مرحلة التحضير وفيه سبع قوائم معبّأة من المحكمين:",
    )
    p_bullet(doc, "المكتسب = مجموع درجات المحكمين في القوائم التابعة لتلك الوحدة والمرحلة.")
    p_bullet(doc, "القصوى = مجموع العلامات المخصصة من معايير التقييم لتلك الوحدة والمرحلة.")
    p_bullet(doc, "نسبة المرحلة = (المكتسب ÷ القصوى) × 100.")
    p_text(doc, "كل مرحلة لها نسبة مستقلة لنفس الوحدة.", before=6)

    heading(doc, "7.2 نتيجة مستوى الوحدة (النهائي للوحدة)", level=2)
    callout(
        doc,
        "المعادلة المعتمدة",
        "نتيجة الوحدة = متوسط نسب مراحلها = (مجموع نسب المراحل ÷ عدد المراحل التي لها نسبة).",
    )
    p_bullet(doc, "ليست متوسطاً مرجّحاً بأوزان مختلفة بين المراحل.")
    p_bullet(doc, "ليست خلطاً بين مستويات وحدات مختلفة؛ كل مستوى وحدة له نتيجته الخاصة.")

    # —— 8 ——
    heading(doc, "8. جدول ملخص سريع")
    add_table(
        doc,
        ["المستوى", "ماذا يُحسب؟"],
        [
            ["القائمة", "نتيجة المحكم داخل قائمة التقييم."],
            [
                "المرحلة × الوحدة",
                "جمع نتائج القوائم ÷ قصوى المعايير لتلك المرحلة والوحدة.",
            ],
            ["نتيجة الوحدة النهائية", "متوسط نسب مراحل تلك الوحدة."],
            [
                "صفحة معايير التقييم",
                "تضبط القصوى والأوزان داخل المرحلة لكل وحدة.",
            ],
            [
                "صفحة التقييم نهائي",
                "تعرض النتيجة لكل مستوى وحدة بعد تجميع مراحله.",
            ],
        ],
    )

    # —— 9 ——
    heading(doc, "9. خلاصة للمراجع")
    p_numbered(
        doc,
        1,
        "آلية معايير التقييم مفعّلة في صفحة التوزيع مع اختيار المرحلة للتفاصيل.",
    )
    p_numbered(
        doc,
        2,
        "عند تعدد القوائم لنفس الوحدة والمرحلة يُحسب مجموع العلامات، ثم تُشتق النسب داخل المرحلة من هذا المجموع.",
    )
    p_numbered(
        doc,
        3,
        "النتيجة النهائية تُبنى أولاً بنسبة كل مرحلة للوحدة، ثم تُلخَّص كنتيجة لمستوى الوحدة بمتوسط نسب مراحلها.",
    )

    p_text(doc, "— نهاية الوثيقة —", size=10, color=MUTED, before=20, after=4)
    p_text(
        doc,
        "يُنصح بمراجعة الأرقام على تمرين تجريبي بعد إدخال علامات مرحلة واحدة على الأقل.",
        size=10,
        color=MUTED,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    DESKTOP.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DESKTOP)
    return OUT


if __name__ == "__main__":
    print(build())
